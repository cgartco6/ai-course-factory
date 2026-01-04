from docx import Document

def create_doc(title, content, filename):
    doc = Document()
    doc.add_heading(title)
    doc.add_paragraph(content)
    doc.save(filename)
    return filename
